from fastapi import APIRouter, UploadFile, File, HTTPException
from fastapi.responses import JSONResponse, FileResponse
import pdfplumber
from docx import Document
from docx.shared import RGBColor
import tempfile
import os
import pandas as pd
from sentence_transformers import SentenceTransformer, util
import torch
import math
import numpy as np
from transformers import pipeline


router = APIRouter(prefix="/redline", tags=["Redlining"])

# ✅ Helper: recursively clean NaN/Inf for safe JSON serialization
def sanitize_for_json(data):
    if isinstance(data, float):
        if math.isnan(data) or math.isinf(data):
            return None
        return data
    elif isinstance(data, dict):
        return {k: sanitize_for_json(v) for k, v in data.items()}
    elif isinstance(data, list):
        return [sanitize_for_json(v) for v in data]
    elif isinstance(data, np.ndarray):
        return sanitize_for_json(data.tolist())
    else:
        return data


# ✅ Load clause playbook safely
try:
    playbook = pd.read_csv("clause_playbook.csv")
except Exception as e:
    raise RuntimeError(f"❌ Could not load clause_playbook.csv: {e}")

required_columns = {"standard_clause", "Risk_Level", "Action_Required"}
if not required_columns.issubset(playbook.columns):
    raise RuntimeError(f"CSV missing required columns: {required_columns - set(playbook.columns)}")

# ✅ Fill blanks
playbook["standard_clause"] = playbook["standard_clause"].fillna("").astype(str)
playbook_clauses = playbook["standard_clause"].tolist()

# ✅ Load LegalBERT model once
model = SentenceTransformer("nlpaueb/legal-bert-base-uncased")

# ✅ Precompute playbook embeddings
playbook_embeddings = model.encode(playbook_clauses, convert_to_tensor=True)

# ✅ Load lightweight AI suggestion model (FLAN-T5)
try:
    suggestion_generator = pipeline("text2text-generation", model="google/flan-t5-base")
except Exception as e:
    raise RuntimeError(f"❌ Failed to load text generation model: {e}")


# ✅ Text extraction functions
def extract_text_from_pdf(file_path):
    text = ""
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text.strip()


def extract_text_from_docx(file_path):
    doc = Document(file_path)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text.strip()


# ✅ Split contract into meaningful clauses
def split_into_clauses(text):
    clauses = [c.strip() for c in text.split(".") if len(c.split()) > 5]
    return clauses if clauses else [text.strip()]


# ✅ Create redlined report as DOCX
def create_redline_doc(results, output_path):
    doc = Document()
    doc.add_heading("Contract Redlining Report", level=1)

    for r in results:
        doc.add_paragraph("📄 Original Clause:", style="List Bullet")
        doc.add_paragraph(r["original_clause"])

        doc.add_paragraph("✅ Matched Standard Clause:")
        doc.add_paragraph(r["matched_clause"])

        doc.add_paragraph(f"🤖 AI-Suggested Clause:")
        doc.add_paragraph(r.get("suggested_clause", "No suggestion available."))

        doc.add_paragraph(f"🔹 Similarity Score: {r['similarity_score']}")

        p = doc.add_paragraph(f"⚠️ Risk Level: {r['risk_level']}")
        run = p.runs[0]
        if r["highlight_color"] == "red":
            run.font.color.rgb = RGBColor(255, 0, 0)
        elif r["highlight_color"] == "orange":
            run.font.color.rgb = RGBColor(255, 165, 0)
        elif r["highlight_color"] == "green":
            run.font.color.rgb = RGBColor(0, 128, 0)

        doc.add_paragraph(f"🛠 Action Required: {r['action_required']}")
        doc.add_paragraph("------------------------------------------------------")

    doc.save(output_path)


# ✅ Main redlining API
@router.post("/analyze/")
async def analyze_contract(file: UploadFile = File(...)):
    try:
        suffix = os.path.splitext(file.filename)[1].lower()
        if suffix not in [".pdf", ".docx"]:
            raise HTTPException(status_code=400, detail="Only PDF or DOCX files are supported.")

        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(await file.read())
            tmp_path = tmp.name

        text = extract_text_from_pdf(tmp_path) if suffix == ".pdf" else extract_text_from_docx(tmp_path)
        os.remove(tmp_path)

        if not text.strip():
            raise HTTPException(status_code=400, detail="Could not extract text from file.")

        clauses = split_into_clauses(text)
        results = []

        for clause in clauses:
            try:
                clause_embedding = model.encode(clause, convert_to_tensor=True)
                scores = util.pytorch_cos_sim(clause_embedding, playbook_embeddings)[0]
                scores = torch.nan_to_num(scores, nan=-1.0, posinf=1.0, neginf=-1.0)

                best_idx = int(torch.argmax(scores))
                best_score = float(scores[best_idx])
                if math.isnan(best_score):
                    best_score = 0.0

                matched_clause = playbook_clauses[best_idx]
                risk_level = str(playbook["Risk_Level"].iloc[best_idx])
                action_required = str(playbook["Action_Required"].iloc[best_idx])

                # ✅ AI-generated suggested clause
                prompt = (
                    f"Rewrite this clause to satisfy the following action requirement:\n\n"
                    f"Clause: {clause}\n"
                    f"Action Required: {action_required}\n\n"
                    f"Provide a legally improved version:"
                )
                try:
                    generated = suggestion_generator(prompt, max_new_tokens=120, num_return_sequences=1)
                    suggested_clause = generated[0]["generated_text"].strip()
                except Exception:
                    suggested_clause = f"Suggestion: {action_required}"

                color = (
                    "green" if risk_level.lower() == "low"
                    else "orange" if risk_level.lower() == "medium"
                    else "red"
                )

                results.append({
                    "original_clause": clause,
                    "matched_clause": matched_clause,
                    "suggested_clause": suggested_clause,
                    "risk_level": risk_level,
                    "action_required": action_required,
                    "similarity_score": round(best_score, 3),
                    "highlight_color": color
                })
            except Exception as inner_err:
                results.append({
                    "original_clause": clause,
                    "error": f"Failed to analyze clause: {inner_err}"
                })

        # ✅ Save DOCX redline report
        output_filename = f"redline_report_{os.path.splitext(file.filename)[0]}.docx"
        output_path = os.path.join("generated_reports", output_filename)
        os.makedirs("generated_reports", exist_ok=True)
        create_redline_doc(results, output_path)

        clean_results = sanitize_for_json({
            "clauses": results,
            "download_link": f"/redline/download/{output_filename}"
        })
        return JSONResponse(content=clean_results)

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ✅ Download endpoint
@router.get("/download/{filename}")
async def download_file(filename: str):
    file_path = os.path.join("generated_reports", filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found.")
    return FileResponse(
        file_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename
    )
