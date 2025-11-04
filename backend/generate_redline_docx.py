from docx import Document
from docx.shared import RGBColor

def generate_redlined_docx(clauses, output_path="redlined_output.docx"):
    """Creates a DOCX file highlighting clauses by risk level."""
    doc = Document()
    doc.add_heading("AI-Generated Redlined Contract Review", level=1)

    for c in clauses:
        clause = c.get("original_clause", "")
        match = c.get("matched_clause", "")
        risk = c.get("risk_level", "")
        action = c.get("action_required", "")
        color = c.get("highlight_color", "black")

        para = doc.add_paragraph()
        run = para.add_run(f"Clause: {clause}\n")
        if color == "red":
            run.font.color.rgb = RGBColor(255, 0, 0)
        elif color == "orange":
            run.font.color.rgb = RGBColor(255, 165, 0)
        elif color == "green":
            run.font.color.rgb = RGBColor(0, 128, 0)

        para.add_run(f"\n→ Matched Standard: {match}")
        para.add_run(f"\n→ Risk Level: {risk}")
        para.add_run(f"\n→ Action: {action}\n")
        doc.add_paragraph("-" * 70)

    doc.save(output_path)
    return output_path
